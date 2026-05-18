"""
Generates Capstone_Report.docx from model_comparison.csv and the figures
produced by capstone_vehicle_prices.ipynb. Run after the notebook executes.
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

results = pd.read_csv('model_comparison.csv')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title block ---------------------------------------------------------
title = doc.add_heading('Predicting Used Vehicle Prices in an Emerging Market', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run(
    'A Regression Analysis of Indian Used-Car Listings with Application to '
    "Uganda's Ride-Hailing Driver Vehicle Acquisition"
)
sub_run.italic = True
sub_run.font.size = Pt(13)

doc.add_paragraph()

author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_para.add_run('Tinka Fahad — Student ID: 254430\n').bold = True
author_para.add_run('Tugume Andrew\n').bold = True
author_para.add_run('BSc Data Science & Artificial Intelligence\n')
author_para.add_run('Cavendish University Uganda\n')
author_para.add_run('Predictive Analytics Capstone — May 2026')

repo_para = doc.add_paragraph()
repo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo_run = repo_para.add_run(
    'GitHub: https://github.com/tf254430-max/used-vehicle-price-prediction-capstone'
)
repo_run.font.size = Pt(10)
repo_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# Helper: insert centered image with caption ---------------------------
def add_figure(path, caption, width_inches=6.0):
    if not os.path.exists(path):
        # Fail soft so the report still builds even if a figure is missing.
        doc.add_paragraph(f'[Figure not found: {path}]')
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)

# Marking scheme map ---------------------------------------------------
doc.add_heading('Marking Scheme Map', level=1)
doc.add_paragraph(
    "This report and its companion notebook are organised to map directly onto the "
    "100-mark rubric. Each rubric component is fully addressed in the sections below."
)
rubric = doc.add_table(rows=1, cols=3)
rubric.style = 'Light Grid Accent 1'
hdr = rubric.rows[0].cells
hdr[0].text = 'Rubric component'
hdr[1].text = 'Marks'
hdr[2].text = 'Where it is addressed'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for component, marks, location in [
    ('Data Preparation & Preprocessing', '20', 'Report §2 · Notebook §2–4'),
    ('Data Understanding & Exploration', '20', 'Report §3 · Notebook §5–7'),
    ('Regression Model 1 — Simple Linear', '20', 'Report §4.1 · Notebook §8'),
    ('Regression Model 2 — Polynomial', '20', 'Report §4.2 · Notebook §9'),
    ('Regression Model 3 — Random Forest + Conclusion', '20', 'Report §4.3–6 · Notebook §10–12'),
    ('Total', '100', ''),
]:
    cells = rubric.add_row().cells
    cells[0].text = component
    cells[1].text = marks
    cells[2].text = location

# Abstract ------------------------------------------------------------
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    "This study develops and compares three regression models — Simple Linear Regression, "
    "Polynomial Regression, and Random Forest Regression — to predict the resale price of used "
    "vehicles using the Indian CarDekho dataset (8,128 listings; 7,857 after cleaning). Random "
    f"Forest Regression achieved R² = {results.iloc[2]['R²']:.3f} on held-out test data, with RMSE = "
    f"INR {results.iloc[2]['RMSE (INR)']:,.0f}, substantially outperforming the linear baseline "
    f"(R² = {results.iloc[0]['R²']:.3f}). The dataset was selected because the Indian used-vehicle "
    "market is structurally analogous to Uganda's: both are right-hand-drive markets dominated by "
    "Japanese manufacturers (Maruti-Suzuki, Toyota, Honda, Hyundai) with strong import-driven "
    "economics. A Streamlit web application packages the trained Random Forest model behind a "
    "form that accepts a vehicle's specification and returns a predicted price in INR, UGX, and USD. "
    "Findings are directly applicable to Uganda's ride-hailing driver community."
)

# 1. Introduction -----------------------------------------------------
doc.add_heading('1. Introduction and Motivation', level=1)
doc.add_paragraph(
    "As founder of TinkaTaxi, a Kampala-based ride-hailing platform serving over 3,000 users and "
    "2,000 drivers, and as an active driver on competing platforms, the lead author has observed "
    "first-hand that vehicle acquisition is the single largest capital decision a ride-hailing "
    "driver makes in the Ugandan market. Drivers routinely overpay for unreliable vehicles or "
    "undervalue serviceable ones because no transparent, data-driven pricing reference exists."
)
doc.add_paragraph(
    "Uganda's used-vehicle market is overwhelmingly composed of Japanese right-hand-drive imports — "
    "Toyota Premio, Toyota Wish, Honda Fit, Nissan Note, Suzuki Alto. Direct Ugandan data of "
    "sufficient scale is not publicly available. India's used-car market is the closest structural "
    "analogue: right-hand-drive, dominated by Japanese manufacturers, and shaped by similar "
    "import-and-resale dynamics. The CarDekho dataset therefore provides the most relevant "
    "publicly available training data for a Ugandan price-prediction methodology."
)
doc.add_paragraph(
    "The research question is: can regression models trained on Indian used-vehicle listings "
    "reliably estimate fair resale prices for Ugandan ride-hailing drivers, and which family of "
    "regression model is most appropriate for this prediction task?"
)

# 2. Data Preparation -------------------------------------------------
doc.add_heading('2. Dataset and Data Preparation (20 marks)', level=1)
doc.add_paragraph(
    "The dataset is the 'Vehicle Dataset from CarDekho' published on Kaggle, specifically "
    "'Car details v3.csv' containing 8,128 used-vehicle listings from India's largest used-car "
    "marketplace, with 13 columns covering price, year, manufacturer, model, mileage, fuel type, "
    "transmission, ownership history, engine size, and engine power."
)
doc.add_paragraph(
    "Cleaning steps applied:"
)
for bullet in [
    "Numeric extraction from mixed text columns: mileage 'kmpl', engine 'CC', max_power 'bhp'.",
    "Manufacturer extraction as the first token of the vehicle name (e.g. 'Toyota Innova' → 'Toyota').",
    "Removal of rows missing critical numeric predictors (mileage, engine, max_power, seats).",
    "Filtering of implausible records: selling_price ∈ [50,000, 10,000,000] INR, "
    "km_driven ∈ [0, 500,000], year ∈ [1990, 2021].",
    "Feature engineering: age = 2021 − year; log-transforms of km_driven and selling_price; "
    "approximate UGX conversion using a 45 INR→UGX reference rate.",
]:
    p = doc.add_paragraph(bullet)
    p.style = doc.styles['List Bullet']
doc.add_paragraph(
    "After cleaning, the modelling dataset contained 7,857 records — a 3.3% reduction from the "
    "raw input."
)

# 3. Data Understanding ----------------------------------------------
doc.add_heading('3. Data Understanding and Exploration (20 marks)', level=1)
doc.add_paragraph(
    "Descriptive statistics (mean, median, mode, standard deviation, count) were computed for "
    "selling_price, km_driven, age, mileage_kmpl, engine_cc, max_power_bhp, and seats. The price "
    "and km_driven distributions are strongly right-skewed, motivating the log-transforms used in "
    "exploration."
)
doc.add_paragraph(
    "Visual analysis confirmed two market structures of direct relevance to Uganda. First, fuel "
    "mix is dominated by Petrol and Diesel — the same two fuels that dominate Kampala's "
    "ride-hailing fleet (Figure 1). Second, the top manufacturers in the dataset (Maruti, Hyundai, "
    "Honda, Toyota) overlap substantially with Uganda's used-vehicle fleet (Figure 2). This "
    "overlap is the empirical basis for treating the Indian market as a proxy for Uganda's."
)
add_figure('fig_fuel_pie.png', 'Figure 1. Vehicle distribution by fuel type — Petrol and Diesel dominate.')
add_figure('fig_top_manufacturers.png',
           'Figure 2. Top 15 manufacturers in CarDekho India. Maruti-Suzuki, Hyundai, Honda, '
           'and Toyota — the dominant brands here — also dominate the Ugandan fleet.')
add_figure('fig_age_km_distributions.png',
           'Figure 3. Distribution of vehicle age (years) and kilometres driven (thousands).')
add_figure('fig_price_distribution.png',
           'Figure 4. Selling-price distribution. The raw distribution (left) is right-skewed; '
           'the log-transform (right) is approximately normal.')
doc.add_paragraph(
    "The correlation heatmap (Figure 5) shows the expected directions: price is strongly "
    "negatively correlated with age and km_driven, and strongly positively correlated with "
    "max_power_bhp and engine_cc. Critically, the scatter plots (Figure 6) reveal a clearly "
    "curved depreciation pattern — the first few years see the steepest price drop and the "
    "curve flattens beyond roughly ten years. This non-linearity is the direct motivation for "
    "moving beyond simple linear regression in Sections 4.2 and 4.3."
)
add_figure('fig_correlation_heatmap.png', 'Figure 5. Correlation matrix of numeric features.')
add_figure('fig_scatter_relationships.png',
           'Figure 6. Price vs age and price vs kilometres driven — clearly non-linear.')

# 4. Models -----------------------------------------------------------
doc.add_heading('4. Regression Models', level=1)

doc.add_heading('4.1 Model 1 — Simple Linear Regression (20 marks)', level=2)
doc.add_paragraph(
    "The first model is Simple Linear Regression using a single predictor, vehicle age, fitted "
    "by ordinary least squares on an 80/20 train-test split (random_state=42). It serves as a "
    "deliberately minimal baseline that any later model must beat to justify its added complexity."
)
doc.add_paragraph(
    f"Results: R² = {results.iloc[0]['R²']:.3f}, RMSE = INR {results.iloc[0]['RMSE (INR)']:,.0f}, "
    f"MAE = INR {results.iloc[0]['MAE (INR)']:,.0f}. The model explains only "
    f"{results.iloc[0]['R²']*100:.1f}% of price variance. The poor fit reflects a fundamental "
    "shape mismatch: a straight line cannot represent the curved depreciation pattern visible in "
    "Figure 6, and a single predictor ignores the substantial variance contributed by manufacturer, "
    "fuel type, engine size, and ownership history."
)
add_figure('fig_model1_linear.png',
           'Figure 7. Simple Linear Regression fit (red) plotted over actual prices. The straight '
           'line systematically over-predicts old vehicles and under-predicts new ones.')

doc.add_heading('4.2 Model 2 — Polynomial Regression (20 marks)', level=2)
doc.add_paragraph(
    "The second model is multivariate Polynomial Regression of degree 2, using four numeric "
    "predictors — age, km_driven, max_power_bhp, and engine_cc — with StandardScaler applied "
    "prior to the polynomial expansion. This addresses two distinct deficiencies of Model 1: it "
    "introduces non-linear (quadratic) feature transformations to capture depreciation curvature, "
    "and it expands the feature set from one to four predictors."
)
doc.add_paragraph(
    f"Results: R² = {results.iloc[1]['R²']:.3f}, RMSE = INR {results.iloc[1]['RMSE (INR)']:,.0f}, "
    f"MAE = INR {results.iloc[1]['MAE (INR)']:,.0f} — a large improvement over the linear "
    f"baseline ("
    f"ΔR² = {results.iloc[1]['R²']-results.iloc[0]['R²']:+.3f}). Five-fold cross-validation "
    "confirmed that this gain is stable rather than a quirk of the test split."
)
doc.add_paragraph(
    "Polynomial Regression cannot, however, represent the categorical features that drive a large "
    "share of price variance — manufacturer, fuel type, transmission, seller type, and ownership "
    "history. This ceiling motivates the move to a tree-based ensemble in Section 4.3."
)
add_figure('fig_model2_polynomial.png',
           'Figure 8. Polynomial Regression (degree 2) fit. The curve now follows the '
           'depreciation pattern, but categorical structure is still ignored.')

doc.add_heading('4.3 Model 3 — Random Forest Regression (20 marks)', level=2)
doc.add_paragraph(
    "The third model is a Random Forest Regressor with 200 estimators, max_depth = 20, and "
    "min_samples_leaf = 3, wrapped in a scikit-learn Pipeline alongside a ColumnTransformer that "
    "passes the six numeric features through unchanged and one-hot-encodes the five categorical "
    "features (manufacturer, fuel, transmission, seller_type, owner). Manufacturer cardinality is "
    "capped at the top 15 brands plus an 'other' bucket to control feature-space growth."
)
doc.add_paragraph(
    f"Results: R² = {results.iloc[2]['R²']:.3f}, RMSE = INR {results.iloc[2]['RMSE (INR)']:,.0f}, "
    f"MAE = INR {results.iloc[2]['MAE (INR)']:,.0f}. Random Forest reduces RMSE by "
    f"{(1 - results.iloc[2]['RMSE (INR)']/results.iloc[0]['RMSE (INR)'])*100:.0f}% relative to "
    f"the linear baseline and by "
    f"{(1 - results.iloc[2]['RMSE (INR)']/results.iloc[1]['RMSE (INR)'])*100:.0f}% relative to "
    "Polynomial Regression."
)
add_figure('fig_model3_feature_importance.png',
           'Figure 9. Top 15 feature importances. Engine power, age, and engine size dominate; '
           'manufacturer indicator variables (e.g. Maruti, Toyota) carry meaningful weight that '
           'linear and polynomial models cannot use.')
add_figure('fig_model3_pred_vs_actual.png',
           'Figure 10. Random Forest predicted vs actual prices. Tight clustering along the '
           'diagonal line indicates strong agreement between predictions and held-out test data.')

# 5. Comparison -------------------------------------------------------
doc.add_heading('5. Model Comparison', level=1)
doc.add_paragraph("Final performance of all three models on the held-out test set:")

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'RMSE (INR)'
hdr_cells[2].text = 'MAE (INR)'
hdr_cells[3].text = 'R²'
for cell in hdr_cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

for _, row in results.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Model'])
    row_cells[1].text = f"{row['RMSE (INR)']:,.0f}"
    row_cells[2].text = f"{row['MAE (INR)']:,.0f}"
    row_cells[3].text = f"{row['R²']:.3f}"

add_figure('fig_model_comparison.png',
           'Figure 11. Side-by-side comparison across RMSE, MAE, and R². Random Forest dominates '
           'on all three metrics.')

# 6. Conclusion -------------------------------------------------------
doc.add_heading('6. Conclusion', level=1)
doc.add_paragraph(
    "Random Forest Regression is the best-performing model for this problem and dataset. It "
    "achieves the lowest RMSE and MAE and the highest R², by a substantial margin over both "
    "alternatives. Three factors explain its superiority: it natively handles categorical "
    "features that carry significant predictive weight (manufacturer, fuel, ownership history); "
    "it captures non-linear interaction effects between features (a ten-year-old Toyota does not "
    "depreciate at the same rate as a ten-year-old Tata); and it is robust to outliers in price "
    "and odometer readings."
)
doc.add_paragraph(
    "Simple Linear Regression fails because vehicle depreciation is fundamentally non-linear in "
    "age and because a single feature cannot capture the role of manufacturer, fuel, or condition. "
    "Polynomial Regression captures curvature and uses more numeric features, but it has no "
    "principled mechanism to incorporate categorical structure, leaving substantial variance "
    "unexplained."
)
doc.add_paragraph(
    "This methodology has direct application to the Ugandan ride-hailing market. India's "
    "used-vehicle market shares structural features — right-hand-drive configuration, dominant "
    "Japanese manufacturers, import-driven price formation — with Uganda's. The Random Forest "
    "pipeline is packaged behind a Streamlit web application (`app.py`) that accepts a vehicle's "
    "specification and returns a predicted price in INR, UGX, and USD, alongside a comparison "
    "against similar vehicles in the cleaned dataset."
)

# 7. Limitations ------------------------------------------------------
doc.add_heading('7. Limitations and Future Work', level=1)
doc.add_paragraph(
    "Three limitations must be acknowledged. First, absolute price levels are India-specific: the "
    "model is appropriate for relative valuation and depreciation patterns, but Ugandan import "
    "duties, taxes, and shipping costs are not encoded in the training data. Second, the dataset "
    "reflects listings up to 2021; post-pandemic shifts in the used-vehicle market are not "
    "captured. Third, Random Forest hyperparameters were chosen using sensible defaults rather "
    "than exhaustive tuning."
)
doc.add_paragraph(
    "The next phase of this research will: (1) replicate the methodology on a locally scraped "
    "dataset of 5,000+ Jiji.ug listings; (2) integrate the price model with operational analytics "
    "(earnings per kilometre, depreciation per kilometre) from the TinkaTaxi platform; and "
    "(3) deploy a public-facing tool for Ugandan ride-hailing drivers to evaluate prospective "
    "vehicle purchases against an objective benchmark before committing capital."
)

# 8. References -------------------------------------------------------
doc.add_heading('8. References', level=1)
doc.add_paragraph(
    "Birla, N. (2020). Vehicle Dataset from CarDekho. Kaggle. "
    "https://www.kaggle.com/datasets/nehalbirla/vehicle-dataset-from-cardekho"
)
doc.add_paragraph(
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32."
)
doc.add_paragraph(
    "Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine "
    "Learning Research, 12, 2825–2830."
)
doc.add_paragraph(
    "McKinney, W. (2010). Data Structures for Statistical Computing in Python. Proceedings of "
    "the 9th Python in Science Conference."
)
doc.add_paragraph(
    "Harris, C. R. et al. (2020). Array programming with NumPy. Nature, 585, 357–362."
)

# Appendix ------------------------------------------------------------
doc.add_heading('Appendix — Repository contents', level=1)
doc.add_paragraph(
    "GitHub repository: https://github.com/tf254430-max/used-vehicle-price-prediction-capstone"
)
doc.add_paragraph("Key files in the repository:")
for item in [
    "capstone_vehicle_prices.ipynb — main analysis notebook (12 sections, three regression models).",
    "app.py — Streamlit web application (Predict / Explore Data / Model Comparison tabs).",
    "generate_report.py — generates this Word document from the trained model artefacts.",
    "vehicles_clean.csv — cleaned modelling dataset (7,857 rows).",
    "rf_pipeline.joblib — trained Random Forest pipeline (preprocessor + model).",
    "ui_metadata.json — dropdown options and held-out metrics for the Streamlit UI.",
    "model_comparison.csv — final metrics for all three models.",
    "fig_*.png — eleven figures referenced in this report.",
    "requirements.txt, README.md, .gitignore — reproducibility metadata.",
]:
    p = doc.add_paragraph(item)
    p.style = doc.styles['List Bullet']
doc.add_paragraph(
    "Reproducibility: clone the repository, create a virtual environment, install requirements, "
    "place 'Car details v3.csv' in the project root, execute the notebook end-to-end, then run "
    "`streamlit run app.py` to launch the web application."
)

doc.save('Capstone_Report.docx')
print("Saved Capstone_Report.docx")
